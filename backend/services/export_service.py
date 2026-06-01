"""Export service — generate Word and PDF reports with desensitization."""
from __future__ import annotations

import base64
import io
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

from backend.models.tag_schema import ReportInstance, ExportConfig
from backend.services.report_generator import _format_chart_value


# ── Desensitization ──────────────────────────────────────

DEFAULT_MASK_RULES = {
    "name": "partial",      # 张三 → 张*
    "phone": "partial",     # 13812345678 → 138****5678
    "id_card": "full",      # → ***
    "address": "partial",   # 保留省市
}


def desensitize(value: str, rule: str = "partial") -> str:
    """Apply desensitization to a string value."""
    if not value or not isinstance(value, str):
        return str(value) if value else ""
    if rule == "full":
        return "***"
    if rule == "partial":
        if len(value) <= 2:
            return value[0] + "*"
        return value[0] + "*" * (len(value) - 2) + value[-1]
    return value


# ── Word export ──────────────────────────────────────────

def export_word(report: ReportInstance, config: ExportConfig | None = None) -> bytes:
    """Generate a Word document from report data."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if config is None:
        config = ExportConfig()

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # ── Title ──
    title = doc.add_heading(f"{report.theme_name}客群画像分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Chapter 1: Summary ──
    doc.add_heading("报告摘要", level=1)
    _add_table(doc, [
        ("报告名称", f"{report.theme_name}客群画像分析报告"),
        ("数据截止日", report.data_date),
        ("生成日期", report.generated_at),
        ("生成用户", report.generated_by),
        ("客群人数", str(report.customer_count)),
    ])

    # ── Chapter 2: Overview ──
    doc.add_heading("一、客群基础概览", level=1)
    if report.overview:
        ov = report.overview
        _add_table(doc, [
            ("客群总人数", str(ov.total_count)),
            ("核心客群", ov.core_segment),
            ("性别分布", _dict_to_str(ov.gender_stats)),
            ("地域分布", _dict_to_str(ov.region_stats, top_n=5)),
            ("资产等级", _dict_to_str(ov.asset_level_stats)),
            ("学历分布", _dict_to_str(ov.education_stats)),
        ])

        if ov.age_histogram:
            doc.add_paragraph("年龄分布：")
            labels = ov.age_histogram.get("labels", [])
            values = ov.age_histogram.get("values", [])
            if labels and values:
                _add_table(doc, list(zip(labels, [str(v) for v in values])))

    # ── Chapter 3: Feature Analysis ──
    doc.add_heading("二、标签特征分析", level=1)
    for group in (report.feature_analysis or []):
        doc.add_heading(group.group_name, level=2)

        for feat in group.features:
            # Feature name + top5 badge
            p = doc.add_paragraph()
            run = p.add_run(f"{'★ ' if feat.is_top5 else ''}{feat.feature_name}")
            run.bold = True
            run.font.size = Pt(11)

            # Feature value
            val_text = _format_chart_value(feat.chart_data, "、")
            if val_text:
                doc.add_paragraph(f"指标值：{val_text}")

            # Chart data as simple table (for non-metric types)
            if feat.chart_type in ("pie", "bar", "histogram"):
                labels = feat.chart_data.get("labels", [])
                values = feat.chart_data.get("values", [])
                if labels and values:
                    _add_table(doc, list(zip(labels, [str(v) for v in values])))

            # LLM insight
            if feat.llm_insight:
                p = doc.add_paragraph()
                run = p.add_run("💡 ")
                run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
                run = p.add_run(feat.llm_insight)
                run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
                run.font.size = Pt(10)

            doc.add_paragraph()  # spacing

    # ── Chapter 4: Correlation Insights ──
    if report.correlation_insights:
        doc.add_heading("三、相关性洞察", level=1)
        for i, rule in enumerate(report.correlation_insights, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"洞察 {i}：")
            run.bold = True

            if rule.type == "enum_enum":
                p.add_run(f"{rule.antecedent} → {rule.consequent}")
                doc.add_paragraph(f"提升度：{rule.lift}，置信度：{rule.confidence:.0%}，支持度：{rule.support:.0%}")
            else:
                p.add_run(f"{rule.antecedent} 的 {rule.consequent} {rule.direction}")
                doc.add_paragraph(f"均值比：{rule.ratio}倍，客群占比：{rule.support:.0%}")

            insight = rule.llm_insight or rule.insight or ""
            if insight:
                p = doc.add_paragraph()
                run = p.add_run("💡 ")
                run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
                p.add_run(insight)

            doc.add_paragraph()

    # ── Chapter 5: Recommendations ──
    if report.recommendations:
        doc.add_heading("四、营销运营建议", level=1)
        rec = report.recommendations

        if rec.marketing_directions:
            doc.add_heading("营销方向", level=2)
            for d in rec.marketing_directions:
                doc.add_paragraph(d, style="List Bullet")

        if rec.priority_customers:
            doc.add_heading("重点跟进客户", level=2)
            doc.add_paragraph(rec.priority_customers)

        if rec.cross_sell_opportunities:
            doc.add_heading("交叉销售机会", level=2)
            for p in rec.cross_sell_opportunities:
                doc.add_paragraph(p, style="List Bullet")

        if rec.marketing_script:
            doc.add_heading("营销话术", level=2)
            p = doc.add_paragraph()
            run = p.add_run("“" + rec.marketing_script + "”")
            run.italic = True
            run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— 报告生成于 {report.generated_at} · IntelliBanker 智能营销平台 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF export (HTML template) ───────────────────────────

def export_pdf(report: ReportInstance, config: ExportConfig | None = None) -> bytes:
    """Generate a PDF from report data using HTML template."""
    from jinja2 import Template

    if config is None:
        config = ExportConfig()

    html = _render_report_html(report.model_dump(mode="json"))

    # Try weasyprint first, fallback to basic HTML
    try:
        from weasyprint import HTML
        return HTML(string=html).write_pdf()
    except ImportError:
        pass

    # Fallback: return HTML as bytes (browser can print to PDF)
    return html.encode("utf-8")


def _render_report_html(report: dict) -> str:
    """Render report as HTML using cached Jinja2 template."""
    return _HTML_TEMPLATE.render(report=report)


# ── HTML template (compiled once at import time) ─────────

_HTML_TEMPLATE_STR = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{{ report.theme_name }}客群画像分析报告</title>
<style>
  body { font-family: "Microsoft YaHei", "PingFang SC", sans-serif; max-width: 900px; margin: 0 auto; padding: 40px; color: #1a212a; line-height: 1.8; font-size: 14px; }
  h1 { text-align: center; color: #1a3a5c; border-bottom: 2px solid #c9a84c; padding-bottom: 12px; }
  h2 { color: #1a3a5c; border-left: 4px solid #c9a84c; padding-left: 12px; margin-top: 32px; }
  h3 { color: #234b73; margin-top: 24px; }
  table { border-collapse: collapse; width: 100%; margin: 12px 0; }
  th, td { border: 1px solid #e5e7eb; padding: 8px 12px; text-align: left; font-size: 13px; }
  th { background: #f8fafc; font-weight: 600; color: #1a3a5c; }
  .metric { display: inline-block; padding: 12px 20px; margin: 6px; background: #f8fafc; border-radius: 8px; text-align: center; min-width: 120px; }
  .metric .value { font-size: 24px; font-weight: 700; color: #1a3a5c; }
  .metric .label { font-size: 12px; color: #6b7280; }
  .insight { padding: 10px 16px; margin: 8px 0; background: #f0f4ff; border-left: 3px solid #c9a84c; font-size: 13px; color: #4a5568; }
  .correlation { padding: 14px 18px; margin: 12px 0; background: #faf5ff; border: 1px solid #e9d5ff; border-radius: 8px; }
  .tag { display: inline-block; padding: 2px 8px; margin: 2px; background: #eff6ff; color: #3b82f6; border-radius: 4px; font-size: 12px; }
  .footer { text-align: center; color: #94a3b8; font-size: 12px; margin-top: 40px; border-top: 1px solid #e5e7eb; padding-top: 16px; }
  @media print { body { padding: 20px; } }
</style>
</head>
<body>
<h1>{{ report.theme_name }}客群画像分析报告</h1>

<h2>报告摘要</h2>
<table>
  <tr><th>报告名称</th><td>{{ report.theme_name }}客群画像分析报告</td></tr>
  <tr><th>数据截止日</th><td>{{ report.data_date }}</td></tr>
  <tr><th>生成日期</th><td>{{ report.generated_at }}</td></tr>
  <tr><th>生成用户</th><td>{{ report.generated_by }}</td></tr>
  <tr><th>客群人数</th><td>{{ report.customer_count }} 人</td></tr>
</table>

{% if report.overview %}
<h2>一、客群基础概览</h2>
<div>
  <div class="metric"><div class="value">{{ report.overview.total_count }}</div><div class="label">客群总人数</div></div>
  <div class="metric"><div class="value">{{ report.overview.core_segment or '—' }}</div><div class="label">核心客群</div></div>
</div>
{% if report.overview.gender_stats %}
<h3>性别分布</h3>
<table><tr>{% for k,v in report.overview.gender_stats.items() %}<th>{{ k }}</th>{% endfor %}</tr>
<tr>{% for k,v in report.overview.gender_stats.items() %}<td>{{ v }}人</td>{% endfor %}</tr></table>
{% endif %}
{% if report.overview.age_histogram and report.overview.age_histogram.labels %}
<h3>年龄分布</h3>
<table><tr>{% for l in report.overview.age_histogram.labels %}<th>{{ l }}</th>{% endfor %}</tr>
<tr>{% for v in report.overview.age_histogram['values'] %}<td>{{ v }}人</td>{% endfor %}</tr></table>
{% endif %}
{% if report.overview.region_stats %}
<h3>地域分布</h3>
<table>{% for k,v in report.overview.region_stats.items() %}<tr><th>{{ k }}</th><td>{{ v }}人</td></tr>{% endfor %}</table>
{% endif %}
{% endif %}

{% for group in report.feature_analysis %}
<h2>{{ loop.index + 1 }}、{{ group.group_name }}</h2>
{% for feat in group.features %}
<h3>{% if feat.is_top5 %}★ {% endif %}{{ feat.feature_name }}</h3>
{% if feat.chart_data.value is defined %}
<p><strong>指标值：</strong>{{ feat.chart_data.value }}{% if feat.chart_data.benchmark %}（全行均值：{{ feat.chart_data.benchmark }}）{% endif %}</p>
{% endif %}
{% if feat.chart_data.get('labels') %}
<table><tr>{% for l in feat.chart_data['labels'] %}<th>{{ l }}</th>{% endfor %}</tr>
<tr>{% for v in feat.chart_data['values'] %}<td>{{ v }}</td>{% endfor %}</tr></table>
{% endif %}
{% if feat.llm_insight %}
<div class="insight">💡 {{ feat.llm_insight }}</div>
{% endif %}
{% endfor %}
{% endfor %}

{% if report.correlation_insights %}
<h2>标签相关性洞察</h2>
{% for rule in report.correlation_insights %}
<div class="correlation">
  <strong>{{ rule.antecedent }} → {{ rule.consequent }}</strong>
  {% if rule.type == 'enum_enum' %}
  <br>提升度：{{ rule.lift }} · 置信度：{{ (rule.confidence * 100)|round(0) }}% · 支持度：{{ (rule.support * 100)|round(0) }}%
  {% else %}
  <br>均值比：{{ rule.ratio }}倍 · {{ rule.direction }}
  {% endif %}
  {% if rule.llm_insight or rule.insight %}
  <div class="insight">💡 {{ rule.llm_insight or rule.insight }}</div>
  {% endif %}
</div>
{% endfor %}
{% endif %}

{% if report.recommendations %}
<h2>营销运营建议</h2>
{% if report.recommendations.marketing_directions %}
<h3>营销方向</h3>
<ul>{% for d in report.recommendations.marketing_directions %}<li>{{ d }}</li>{% endfor %}</ul>
{% endif %}
{% if report.recommendations.priority_customers %}
<h3>重点跟进客户</h3>
<p>{{ report.recommendations.priority_customers }}</p>
{% endif %}
{% if report.recommendations.cross_sell_opportunities %}
<h3>交叉销售机会</h3>
<ul>{% for p in report.recommendations.cross_sell_opportunities %}<li>{{ p }}</li>{% endfor %}</ul>
{% endif %}
{% if report.recommendations.marketing_script %}
<h3>营销话术</h3>
<p style="font-style:italic; color:#4a5568; padding:12px 16px; background:#f0f4ff; border-radius:8px;">"{{ report.recommendations.marketing_script }}"</p>
{% endif %}
{% endif %}

<div class="footer">— 报告生成于 {{ report.generated_at }} · IntelliBanker 智能营销平台 —</div>
</body></html>"""

from jinja2 import Template as _JinjaTemplate
_HTML_TEMPLATE = _JinjaTemplate(_HTML_TEMPLATE_STR)


# ── Helpers ──────────────────────────────────────────────

def _add_table(doc, rows: list[tuple[str, str]]):
    """Add a two-column table to the document."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        # Bold key column
        for p in table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True


def _dict_to_str(d: dict, top_n: int = 0) -> str:
    """Convert dict to display string."""
    if not d:
        return "—"
    items = list(d.items())
    if top_n > 0:
        items = items[:top_n]
    return "、".join(f"{k}({v})" for k, v in items)
